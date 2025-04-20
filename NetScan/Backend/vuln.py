import streamlit as st
import socket
import requests
import subprocess
import re
import concurrent.futures
import time
import urllib.parse
from bs4 import BeautifulSoup
import random
import json
import base64
import io
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="IP Vulnerability Detector", layout="wide")
st.title("🛡️ IP Vulnerability Detector")
st.markdown("Enter an IP address to scan for open ports and search for known vulnerabilities.")

# You can either use an API key or leave as empty string to use fallback method
SERPAPI_KEY = ""  # Your SerpAPI key if you have one, otherwise leave empty


# Optimized port scanner using threading for parallel scanning
def scan_port(ip, port):
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.settimeout(0.1)  # Reduced timeout for faster scanning
    result = sock.connect_ex((ip, port))
    sock.close()
    if result == 0:
        try:
            service = socket.getservbyport(port)
        except:
            service = "unknown"
        return port, service
    return None


def scan_ports(ip, port_range=None):
    if port_range is None:
        port_range = range(1, 1025)

    open_ports = []
    # Use ThreadPoolExecutor for parallel port scanning
    with concurrent.futures.ThreadPoolExecutor(max_workers=100) as executor:
        # Submit all port scan tasks
        future_to_port = {executor.submit(scan_port, ip, port): port for port in port_range}

        # Process results as they complete
        for future in concurrent.futures.as_completed(future_to_port):
            result = future.result()
            if result:
                open_ports.append(result)

    return open_ports


# Optimized nmap service detection - scan only for common ports if too many are open
def detect_service_versions(ip, open_ports):
    services = []
    try:
        # Limit to 20 ports max for speed
        ports_to_scan = open_ports[:20] if len(open_ports) > 20 else open_ports
        port_list = ",".join(str(p[0]) for p in ports_to_scan)

        # Run nmap with faster options
        result = subprocess.run(
            ["nmap", "-T4", "-sV", "--version-intensity", "0", "-p", port_list, ip],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True,
            timeout=15  # Add timeout to prevent hanging
        )
        output = result.stdout

        # Format and display nmap output in a more readable way
        display_formatted_nmap_output(output, ip)

        for port, _ in ports_to_scan:
            # Look for the line containing this port
            # Improved regex to handle more nmap output formats
            line_match = re.search(rf"{port}/tcp\s+open\s+(\S+(?:/\S+)?)(?:\s+(.*))?", output)

            if line_match:
                service = line_match.group(1)  # Service name
                version_info = line_match.group(2) if line_match.lastindex == 2 else ""

                # Extract version from parenthetical info if present
                # Look for version pattern like X.Y.Z
                version = re.search(r"(\d+(\.\d+)+)", version_info)
                version_str = version.group(1) if version else "unknown"

                # For services like "ssl/domain", extract the main part
                if "/" in service:
                    service_parts = service.split("/")
                    # Use the second part as the primary service
                    service = service_parts[1] if len(service_parts) > 1 else service_parts[0]
            else:
                service = "unknown"
                version_str = "unknown"

            services.append({"port": port, "service": service, "version": version_str})

        # For remaining ports, add with basic info
        if len(open_ports) > 20:
            for port, service in open_ports[20:]:
                services.append({"port": port, "service": service, "version": "unknown"})
    except Exception as e:
        st.warning(f"Error detecting services: {e}")
        for port, service in open_ports:
            services.append({"port": port, "service": service, "version": "unknown"})
    return services


# New function to format and display nmap output in a structured way
def display_formatted_nmap_output(output, ip):
    # Create an expandable section for the raw nmap output
    with st.expander("Nmap Scan Details", expanded=False):
        # Extract key information using regex
        start_time_match = re.search(r"Starting Nmap .* at (.*)", output)
        start_time = start_time_match.group(1) if start_time_match else "Unknown"

        latency_match = re.search(r"Host is up \((.*?)\)", output)
        latency = latency_match.group(1) if latency_match else "Unknown"

        scan_time_match = re.search(r"scanned in (.*)", output)
        scan_time = scan_time_match.group(1) if scan_time_match else "Unknown"

        # Format port information
        port_info = re.findall(r"(\d+)/tcp\s+open\s+(\S+(?:/\S+)?)(?:\s+(.*))?", output)

        # Create a simple markdown table for port information
        if port_info:
            st.subheader("Open Ports")
            st.markdown("| Port | Service | Version |")
            st.markdown("|------|---------|---------|")
            for port, service, version in port_info:
                version_clean = version if version else "Unknown"
                st.markdown(f"| {port} | {service} | {version_clean} |")

        # Display scan metadata
        st.subheader("Scan Details")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"**Target IP:** {ip}")
        with col2:
            st.markdown(f"**Latency:** {latency}")
        with col3:
            st.markdown(f"**Scan Duration:** {scan_time}")

        st.markdown(f"**Scan Started:** {start_time}")

        # Show the raw output in a code block for reference
        st.subheader("Raw Nmap Output")
        st.code(output)


# Google search for vulnerabilities
def search_google_for_vulnerabilities(query):
    """Search Google for vulnerabilities related to the service and parse results."""
    # Log that the Google search is being attempted
    st.info(f"No results found in NVD database. Searching Google for: {query}")

    search_query = f"{query} vulnerability security exploit"
    encoded_query = urllib.parse.quote_plus(search_query)

    # Using a different search API if key is available
    if SERPAPI_KEY:
        url = f"https://serpapi.com/search.json?q={encoded_query}&api_key={SERPAPI_KEY}"

        try:
            response = requests.get(url, timeout=10)
            data = response.json()

            results = []
            if "organic_results" in data:
                for i, result in enumerate(data["organic_results"][:5]):
                    title = result.get("title", "Unknown Title")
                    link = result.get("link", "#")
                    snippet = result.get("snippet", "No description available")

                    # Extract any CVE IDs from title or description
                    cve_pattern = r'CVE-\d{4}-\d{4,}'
                    cve_matches = re.findall(cve_pattern, title + " " + snippet)
                    cve_id = cve_matches[0] if cve_matches else f"GOOGLE-RESULT-{i + 1}"

                    # Try to determine severity from keywords
                    severity = "Unknown"
                    severity_score = "N/A"

                    severity_keywords = {
                        "CRITICAL": ["critical", "severe", "urgent", "emergency"],
                        "HIGH": ["high", "important", "major"],
                        "MEDIUM": ["medium", "moderate"],
                        "LOW": ["low", "minor"]
                    }

                    content = (title + " " + snippet).lower()
                    for sev, keywords in severity_keywords.items():
                        if any(keyword in content for keyword in keywords):
                            severity = sev
                            break

                    results.append({
                        "id": cve_id,
                        "description": snippet,
                        "published": "Unknown",
                        "severity": severity,
                        "score": severity_score,
                        "source": "Google Search",
                        "title": title,
                        "link": link
                    })

            if results:
                return results
        except Exception as e:
            st.error(f"Google search API error: {e}")

    # Fallback to direct scraping
    return fallback_google_search(query)


def fallback_google_search(query):
    """Fallback method that attempts to scrape Google directly."""
    search_query = f"{query} vulnerability security exploit CVE"
    encoded_query = urllib.parse.quote_plus(search_query)
    url = f"https://www.google.com/search?q={encoded_query}"

    # Rotate between different user agents to avoid detection
    user_agents = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.0 Safari/605.1.15",
        "Mozilla/5.0 (X11; Linux x86_64; rv:89.0) Gecko/20100101 Firefox/89.0"
    ]

    headers = {
        "User-Agent": random.choice(user_agents),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "Referer": "https://www.google.com/",
        "DNT": "1",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1"
    }

    try:
        # Add delay to avoid rate limiting
        time.sleep(0.5)
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')

        results = []
        # Try different CSS selectors that Google might use
        search_divs = soup.find_all('div', attrs={'class': ['g', 'tF2Cxc', 'yuRUbf']})

        if not search_divs:
            # Try another common pattern
            search_divs = soup.select("div.g")

        for i, div in enumerate(search_divs[:5]):  # Limit to first 5 results
            try:
                # Try different patterns to extract title, link, and description
                title_elem = div.find('h3') or div.select_one("h3")
                link_elem = div.find('a') or div.select_one("a")

                # Different ways to find description
                desc_elem = div.find('div', class_=['VwiC3b', 'yXK7lf', 'MUxGbd', 'lyLwlc']) or \
                            div.select_one("span.st") or \
                            div.select_one("div.IsZvec")

                if title_elem and link_elem:
                    title = title_elem.get_text()
                    link = link_elem.get('href')

                    desc = ""
                    if desc_elem:
                        desc = desc_elem.get_text()

                    # Extract any CVE IDs from title or description
                    cve_pattern = r'CVE-\d{4}-\d{4,}'
                    cve_matches = re.findall(cve_pattern, title + " " + desc)
                    cve_id = cve_matches[0] if cve_matches else f"GOOGLE-RESULT-{i + 1}"

                    # Try to determine severity from keywords
                    severity = "Unknown"
                    severity_score = "N/A"

                    severity_keywords = {
                        "CRITICAL": ["critical", "severe", "urgent", "emergency"],
                        "HIGH": ["high", "important", "major"],
                        "MEDIUM": ["medium", "moderate"],
                        "LOW": ["low", "minor"]
                    }

                    content = (title + " " + desc).lower()
                    for sev, keywords in severity_keywords.items():
                        if any(keyword in content for keyword in keywords):
                            severity = sev
                            break

                    results.append({
                        "id": cve_id,
                        "description": desc or "No description available",
                        "published": "Unknown",
                        "severity": severity,
                        "score": severity_score,
                        "source": "Google Search",
                        "title": title,
                        "link": link
                    })
            except Exception as e:
                st.error(f"Error parsing search result: {e}")
                continue

        # If still no results, try a simpler approach
        if not results:
            # Create dummy results based on service name
            results.append({
                "id": "POTENTIAL-VULNERABILITY",
                "description": f"Consider further manual investigation for security issues with {query}",
                "published": "Unknown",
                "severity": "Unknown",
                "score": "N/A",
                "source": "Recommendation",
                "title": f"Potential security considerations for {query}",
                "link": f"https://www.google.com/search?q={encoded_query}"
            })

        return results
    except Exception as e:
        st.error(f"Fallback Google search error: {e}")
        return [{
            "id": "SEARCH-FAILED",
            "description": f"Google search failed, but you should manually investigate {query} for potential vulnerabilities",
            "published": "Unknown",
            "severity": "Unknown",
            "score": "N/A",
            "source": "Search Failed",
            "title": f"Manual investigation needed for {query}",
            "link": f"https://www.google.com/search?q={encoded_query}"
        }]


# Cache for previously fetched vulnerabilities to avoid duplicate API calls
vulnerability_cache = {}


def fetch_vulnerabilities(service_name, version=""):
    if service_name == "unknown":
        # Even for unknown services, try a generic search
        return search_google_for_vulnerabilities("port security vulnerabilities")

    query = service_name
    if version and version != "unknown":
        query += f" {version}"

    # Check cache first
    if query in vulnerability_cache:
        return vulnerability_cache[query]

    url = "https://services.nvd.nist.gov/rest/json/cves/2.0"
    params = {"keywordSearch": query, "resultsPerPage": 5}

    try:
        response = requests.get(url, params=params)
        data = response.json()

        vulns = []
        if "vulnerabilities" in data and data["vulnerabilities"]:
            for vuln in data["vulnerabilities"]:
                cve = vuln["cve"]
                vuln_info = {
                    "id": cve["id"],
                    "description": cve["descriptions"][0]["value"],
                    "published": cve.get("published", ""),
                    "severity": "Unknown",
                    "score": "N/A",
                    "source": "NVD"
                }
                metrics = cve.get("metrics", {})
                if "cvssMetricV31" in metrics:
                    cvss = metrics["cvssMetricV31"][0]["cvssData"]
                    vuln_info["severity"] = cvss["baseSeverity"]
                    vuln_info["score"] = cvss["baseScore"]
                elif "cvssMetricV2" in metrics:
                    cvss = metrics["cvssMetricV2"][0]["cvssData"]
                    score = cvss["baseScore"]
                    vuln_info["score"] = score
                    if score >= 7.0:
                        vuln_info["severity"] = "HIGH"
                    elif score >= 4.0:
                        vuln_info["severity"] = "MEDIUM"
                    else:
                        vuln_info["severity"] = "LOW"

                vulns.append(vuln_info)

        # If no vulnerabilities found in NVD, search Google
        if not vulns:
            st.info(f"No vulnerabilities found in NVD for {query}. Searching Google...")
            google_vulns = search_google_for_vulnerabilities(query)
            vulns.extend(google_vulns)

        # Store in cache
        vulnerability_cache[query] = vulns
        return vulns
    except Exception as e:
        st.error(f"Failed to fetch vulnerabilities from NVD: {e}")
        # Try Google search as fallback
        try:
            st.warning(f"NVD search failed. Using Google search as fallback for {query}...")
            google_vulns = search_google_for_vulnerabilities(query)
            vulnerability_cache[query] = google_vulns
            return google_vulns
        except Exception as e2:
            st.error(f"Failed to fetch vulnerabilities from Google: {e2}")
            return [{
                "id": "SEARCH-FAILED",
                "description": f"All vulnerability searches failed. Please manually investigate {query} for potential security issues.",
                "published": "Unknown",
                "severity": "Unknown",
                "score": "N/A",
                "source": "Search Failed",
                "title": f"Manual investigation needed for {query}",
                "link": f"https://www.google.com/search?q={urllib.parse.quote_plus(query + ' vulnerability security')}"
            }]


# Add a function to create severity badges with appropriate colors
def get_severity_badge(severity):
    severity_colors = {
        "CRITICAL": "red",
        "HIGH": "orange",
        "MEDIUM": "yellow",
        "LOW": "blue",
        "Unknown": "gray"
    }
    color = severity_colors.get(severity, "gray")
    return f"<span style='background-color:{color}; color:white; padding:2px 6px; border-radius:3px; font-size:0.8em'>{severity}</span>"


# Export functions
def get_download_link(content, filename, text):
    """Generate a download link for a file"""
    b64 = base64.b64encode(content).decode()
    href = f'<a href="data:file/txt;base64,{b64}" download="{filename}">{text}</a>'
    return href


def export_to_txt(ip_address, services, service_vulnerabilities):
    """Export scan results to TXT format"""
    output = io.StringIO()

    # Header
    output.write(f"IP Vulnerability Scan Results\n")
    output.write(f"Target: {ip_address}\n")
    output.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    output.write(f"{'-' * 50}\n\n")

    # Summary
    output.write(f"SUMMARY:\n")
    output.write(f"Found {len(services)} open ports\n\n")

    # Services and vulnerabilities
    for service_info in services:
        port = service_info['port']
        service = service_info['service']
        version = service_info['version']
        vulns = service_vulnerabilities.get(port, [])

        output.write(f"\nPORT {port} - {service} {version if version != 'unknown' else ''}\n")
        output.write(f"{'-' * 50}\n")

        if vulns:
            output.write(f"Found {len(vulns)} potential vulnerabilities\n\n")

            for vuln in vulns:
                output.write(f"ID: {vuln.get('id', 'Unknown')}\n")
                output.write(f"Severity: {vuln.get('severity', 'Unknown')} (Score: {vuln.get('score', 'N/A')})\n")
                output.write(f"Title: {vuln.get('title', 'Unknown')}\n")
                output.write(f"Description: {vuln.get('description', 'No description available')}\n")
                output.write(f"Source: {vuln.get('source', 'Unknown')}\n")
                if vuln.get('link'):
                    output.write(f"Link: {vuln.get('link')}\n")
                output.write(f"{'-' * 40}\n\n")
        else:
            output.write(f"No known vulnerabilities found for this service.\n\n")

    return output.getvalue().encode()


def export_to_json(ip_address, services, service_vulnerabilities):
    """Export scan results to JSON format"""
    data = {
        "scan_info": {
            "target_ip": ip_address,
            "scan_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "open_ports_count": len(services)
        },
        "services": []
    }

    for service_info in services:
        port = service_info['port']
        service = service_info['service']
        version = service_info['version']
        vulns = service_vulnerabilities.get(port, [])

        service_data = {
            "port": port,
            "service": service,
            "version": version,
            "vulnerabilities": vulns
        }

        data["services"].append(service_data)

    return json.dumps(data, indent=4).encode()


def export_to_pdf(ip_address, services, service_vulnerabilities):
    """Export scan results to PDF format"""
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    # Add title style
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=1,  # Center alignment
        spaceAfter=12
    )

    # Add heading style
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6
    )

    # Add subheading style
    subheading_style = ParagraphStyle(
        'Subheading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceBefore=6,
        spaceAfter=3
    )

    # Add text style
    text_style = styles['Normal']

    # Create the document structure
    elements = []

    # Title
    elements.append(Paragraph(f"IP Vulnerability Scan Results", title_style))
    elements.append(Spacer(1, 12))

    # Scan Information
    elements.append(Paragraph("Scan Information", heading_style))
    scan_info = [
        ["Target IP:", ip_address],
        ["Scan Date:", datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ["Open Ports:", str(len(services))]
    ]
    scan_info_table = Table(scan_info, colWidths=[120, 350])
    scan_info_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 6)
    ]))
    elements.append(scan_info_table)
    elements.append(Spacer(1, 12))

    # Services and vulnerabilities
    for service_info in services:
        port = service_info['port']
        service = service_info['service']
        version = service_info['version']
        vulns = service_vulnerabilities.get(port, [])

        # Service heading
        elements.append(Paragraph(f"Port {port} - {service} {version if version != 'unknown' else ''}", heading_style))
        elements.append(Spacer(1, 6))

        if vulns:
            elements.append(Paragraph(f"Found {len(vulns)} potential vulnerabilities", text_style))
            elements.append(Spacer(1, 6))

            # Group vulnerabilities by source
            nvd_vulns = [v for v in vulns if v.get("source") == "NVD"]
            google_vulns = [v for v in vulns if v.get("source") == "Google Search"]
            other_vulns = [v for v in vulns if v.get("source") not in ["NVD", "Google Search"]]

            # Process and display vulnerabilities from different sources
            all_vuln_groups = [
                ("NVD Database", nvd_vulns),
                ("Google Results", google_vulns),
                ("Other Sources", other_vulns)
            ]

            for source_name, source_vulns in all_vuln_groups:
                if source_vulns:
                    elements.append(Paragraph(source_name, subheading_style))

                    for vuln in source_vulns:
                        # Create a vulnerability info table
                        vuln_data = [
                            ["ID:", vuln.get('id', 'Unknown')],
                            ["Severity:", f"{vuln.get('severity', 'Unknown')} (Score: {vuln.get('score', 'N/A')})"],
                            ["Description:", vuln.get('description', 'No description available')],
                        ]

                        if vuln.get('title'):
                            vuln_data.insert(1, ["Title:", vuln.get('title')])

                        if vuln.get('published'):
                            vuln_data.append(["Published:", vuln.get('published', 'Unknown')])

                        if vuln.get('link'):
                            vuln_data.append(["Link:", vuln.get('link')])

                        vuln_table = Table(vuln_data, colWidths=[80, 390])
                        vuln_table.setStyle(TableStyle([
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('PADDING', (0, 0), (-1, -1), 4)
                        ]))
                        elements.append(vuln_table)
                        elements.append(Spacer(1, 10))
        else:
            elements.append(Paragraph("No known vulnerabilities found for this service.", text_style))

        elements.append(Spacer(1, 12))

    # Build the PDF
    pdf.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_docx(ip_address, services, service_vulnerabilities):
    """Export scan results to DOCX format"""
    doc = Document()

    # Add title
    doc.add_heading('IP Vulnerability Scan Results', 0)

    # Add scan information
    doc.add_heading('Scan Information', level=1)
    scan_info_table = doc.add_table(rows=3, cols=2)
    scan_info_table.style = 'Table Grid'

    # Fill scan info table
    cells = scan_info_table.rows[0].cells
    cells[0].text = 'Target IP:'
    cells[1].text = ip_address

    cells = scan_info_table.rows[1].cells
    cells[0].text = 'Scan Date:'
    cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    cells = scan_info_table.rows[2].cells
    cells[0].text = 'Open Ports:'
    cells[1].text = str(len(services))

    # Apply formatting to scan info table
    for row in scan_info_table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            cell.width = Inches(2)

    doc.add_paragraph('')

    # Add services and vulnerabilities
    for service_info in services:
        port = service_info['port']
        service = service_info['service']
        version = service_info['version']
        vulns = service_vulnerabilities.get(port, [])

        # Add service heading
        service_heading = doc.add_heading(f'Port {port} - {service} {version if version != "unknown" else ""}', level=1)

        if vulns:
            doc.add_paragraph(f'Found {len(vulns)} potential vulnerabilities')

            # Group vulnerabilities by source
            nvd_vulns = [v for v in vulns if v.get("source") == "NVD"]
            google_vulns = [v for v in vulns if v.get("source") == "Google Search"]
            other_vulns = [v for v in vulns if v.get("source") not in ["NVD", "Google Search"]]

            # Process and display vulnerabilities from different sources
            all_vuln_groups = [
                ("NVD Database", nvd_vulns),
                ("Google Results", google_vulns),
                ("Other Sources", other_vulns)
            ]

            for source_name, source_vulns in all_vuln_groups:
                if source_vulns:
                    doc.add_heading(source_name, level=2)

                    for vuln in source_vulns:
                        # Create paragraph for vulnerability title
                        title_text = vuln.get('title', vuln.get('id', 'Unknown Vulnerability'))
                        title_para = doc.add_paragraph()
                        title_run = title_para.add_run(title_text)
                        title_run.bold = True
                        title_run.font.size = Pt(12)

                        # Add severity info
                        severity = vuln.get('severity', 'Unknown')
                        score = vuln.get('score', 'N/A')
                        severity_para = doc.add_paragraph()
                        severity_run = severity_para.add_run(f"Severity: {severity} (Score: {score})")

                        # Color-code severity
                        if severity == "CRITICAL":
                            severity_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                        elif severity == "HIGH":
                            severity_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                        elif severity == "MEDIUM":
                            severity_run.font.color.rgb = RGBColor(255, 255, 0)  # Yellow
                        elif severity == "LOW":
                            severity_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue

                        # Add description
                        doc.add_paragraph(f"Description: {vuln.get('description', 'No description available')}")

                        # Add published date if available
                        if vuln.get('published'):
                            doc.add_paragraph(f"Published: {vuln.get('published')}")

                        # Add link if available
                        if vuln.get('link'):
                            link_para = doc.add_paragraph("Link: ")
                            link_run = link_para.add_run(vuln.get('link'))
                            link_run.font.underline = True
                            link_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for links

                        doc.add_paragraph('─' * 40)  # Add separator
        else:
            no_vuln_para = doc.add_paragraph("No known vulnerabilities found for this service.")
            no_vuln_para.italic = True

        doc.add_paragraph('')  # Add space between services

    # Save to a BytesIO object
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# Main UI components and logic
st.sidebar.header("Scan Options")
ip_address = st.sidebar.text_input("Target IP Address", "8.8.8.8")

# Add fast scan option
scan_speed = st.sidebar.radio(
    "Scan Speed",
    ["Fast (Common Ports)", "Normal (1000 ports)", "Thorough (Custom Range)"]
)

if scan_speed == "Fast (Common Ports)":
    # Common ports to scan for speed
    common_ports = [20, 21, 22, 23, 25, 53, 80, 110, 143, 443, 465, 587, 853, 993, 995, 3306, 3389, 5432, 5900, 8080,
                    8443]
    port_range = (min(common_ports), max(common_ports))
    st.sidebar.info(f"Scanning common ports: {', '.join(map(str, common_ports))}")
elif scan_speed == "Normal (1000 ports)":
    port_range = (1, 1000)
else:
    port_range = st.sidebar.slider("Port Range", 1, 10000, (1, 1024))

scan_button = st.sidebar.button("Start Scan")

# Add progress indicator
progress_placeholder = st.empty()

tab1, tab2 = st.tabs(["Scan Results", "About"])

with tab1:
    if scan_button:
        start_time = time.time()

        # Initialize progress
        progress_bar = progress_placeholder.progress(0)

        # Define ports to scan based on selection
        if scan_speed == "Fast (Common Ports)":
            ports_to_scan = common_ports
        else:
            ports_to_scan = range(port_range[0], port_range[1] + 1)

        with st.spinner("Scanning ports..."):
            progress_bar.progress(10)
            open_ports = scan_ports(ip_address, ports_to_scan)
            progress_bar.progress(50)

        if open_ports:
            # Sort ports by number for better readability
            open_ports.sort(key=lambda x: x[0])

            st.success(f"Found {len(open_ports)} open ports on {ip_address}")

            # Create a simple markdown table for open ports
            with st.spinner("Detecting service versions..."):
                progress_bar.progress(60)
                services = detect_service_versions(ip_address, open_ports)
                progress_bar.progress(75)

            # Process vulnerabilities in parallel
            service_vulnerabilities = {}

            # First, submit all vulnerability fetch tasks
            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                future_to_service = {
                    executor.submit(fetch_vulnerabilities, service_info['service'], service_info['version']):
                        service_info for service_info in services
                }

                # Process results as they complete
                completed = 0
                total = len(future_to_service)
                for future in concurrent.futures.as_completed(future_to_service):
                    service_info = future_to_service[future]
                    service_vulnerabilities[service_info['port']] = future.result()
                    completed += 1
                    progress_bar.progress(75 + (25 * completed // total))

            progress_bar.progress(100)

            # Display services and vulnerabilities using cards
            st.subheader("Vulnerability Analysis by Port")

            for service_info in services:
                port = service_info['port']
                service = service_info['service']
                version = service_info['version']
                vulns = service_vulnerabilities.get(port, [])

                # Create a card-like container for each port
                with st.container():
                    st.markdown(f"### Port {port} - {service} {version if version != 'unknown' else ''}")
                    st.markdown("---")

                    if vulns:
                        # Count sources
                        nvd_vulns = [v for v in vulns if v.get("source") == "NVD"]
                        google_vulns = [v for v in vulns if v.get("source") == "Google Search"]
                        other_vulns = [v for v in vulns if v.get("source") not in ["NVD", "Google Search"]]

                        # Show summary
                        sources_text = []
                        if nvd_vulns:
                            sources_text.append(f"{len(nvd_vulns)} from NVD")
                        if google_vulns:
                            sources_text.append(f"{len(google_vulns)} from Google")
                        if other_vulns:
                            sources_text.append(f"{len(other_vulns)} from other sources")

                        st.success(f"Found {len(vulns)} potential vulnerabilities ({', '.join(sources_text)})")

                        # Create tabs for different vulnerability sources
                        vuln_tabs = []
                        if nvd_vulns:
                            vuln_tabs.append("NVD Database")
                        if google_vulns:
                            vuln_tabs.append("Google Results")
                        if other_vulns:
                            vuln_tabs.append("Other Sources")

                        if vuln_tabs:
                            selected_tab = st.radio(f"Vulnerability sources for Port {port}", vuln_tabs,
                                                    key=f"tabs_{port}")

                            if selected_tab == "NVD Database" and nvd_vulns:
                                for vuln in nvd_vulns:
                                    severity = vuln.get("severity", "Unknown")
                                    severity_badge = get_severity_badge(severity)

                                    st.markdown(f"""
                                    <div style="border:1px solid #eee; border-radius:5px; padding:10px; margin-bottom:10px">
                                        <h4>{vuln['id']} {severity_badge} (Score: {vuln.get('score', 'N/A')})</h4>
                                        <p>{vuln["description"]}</p>
                                        <small>Published: {vuln.get('published', 'Unknown')}</small>
                                    </div>
                                    """, unsafe_allow_html=True)

                            elif selected_tab == "Google Results" and google_vulns:
                                for vuln in google_vulns:
                                    severity = vuln.get("severity", "Unknown")
                                    severity_badge = get_severity_badge(severity)

                                    st.markdown(f"""
                                    <div style="border:1px solid #eee; border-radius:5px; padding:10px; margin-bottom:10px">
                                        <h4>{vuln['title']} {severity_badge}</h4>
                                        <p>{vuln["description"]}</p>
                                        <a href="{vuln.get('link', '#')}" target="_blank">View Details</a>
                                    </div>
                                    """, unsafe_allow_html=True)

                            elif selected_tab == "Other Sources" and other_vulns:
                                for vuln in other_vulns:
                                    st.markdown(f"""
                                    <div style="border:1px solid #eee; border-radius:5px; padding:10px; margin-bottom:10px">
                                        <h4>{vuln.get('title', vuln.get('id', 'Unknown'))}</h4>
                                        <p>{vuln["description"]}</p>
                                        <a href="{vuln.get('link', '#')}" target="_blank">View Details</a>
                                    </div>
                                    """, unsafe_allow_html=True)
                    else:
                        # If no vulnerabilities found, show a message with a manual search link
                        search_query = f"{service} {version if version != 'unknown' else ''} vulnerability security"
                        encoded_query = urllib.parse.quote_plus(search_query)
                        google_link = f"https://www.google.com/search?q={encoded_query}"

                        st.info(f"""No known vulnerabilities found for this service.

You can [search Google manually]({google_link}) for potential security information.""")

            # Display scan time
            end_time = time.time()
            st.success(f"Scan completed in {end_time - start_time:.2f} seconds")

            # Add export functionality
            st.markdown("---")
            st.subheader("Export Results")

            col1, col2 = st.columns(2)

            with col1:
                st.markdown("#### Export Format")
                export_format = st.selectbox("Choose export format:",
                                             ["PDF", "JSON", "DOCX", "TXT"])

            with col2:
                st.markdown("#### Download")
                if export_format == "PDF":
                    pdf_data = export_to_pdf(ip_address, services, service_vulnerabilities)
                    st.markdown(get_download_link(pdf_data, f"vulnerability_scan_{ip_address}.pdf",
                                                  "📥 Download PDF Report"), unsafe_allow_html=True)
                elif export_format == "JSON":
                    json_data = export_to_json(ip_address, services, service_vulnerabilities)
                    st.markdown(get_download_link(json_data, f"vulnerability_scan_{ip_address}.json",
                                                  "📥 Download JSON Data"), unsafe_allow_html=True)
                elif export_format == "DOCX":
                    docx_data = export_to_docx(ip_address, services, service_vulnerabilities)
                    st.markdown(get_download_link(docx_data, f"vulnerability_scan_{ip_address}.docx",
                                                  "📥 Download DOCX Report"), unsafe_allow_html=True)
                elif export_format == "TXT":
                    txt_data = export_to_txt(ip_address, services, service_vulnerabilities)
                    st.markdown(get_download_link(txt_data, f"vulnerability_scan_{ip_address}.txt",
                                                  "📥 Download TXT Report"), unsafe_allow_html=True)
        else:
            progress_bar.progress(100)
            st.warning(f"No open ports found on {ip_address} in the specified range.")

with tab2:
    st.header("About this Tool")
    st.write("""
    This vulnerability scanner works by:

    1. Scanning open ports on the target IP
    2. Identifying services and versions running on those ports
    3. Querying the National Vulnerability Database (NVD) for known vulnerabilities
    4. If no vulnerabilities are found in NVD, searching Google for relevant security information

    ### Features:
    - Parallel port scanning for faster results
    - Optimized nmap detection
    - In-memory vulnerability caching
    - Multiple scan speed options
    - Real-time CVE lookup from NVD API
    - Google search fallback for services not in NVD
    - Formatted and readable scan results
    - Export reports in multiple formats (PDF, JSON, DOCX, TXT)

    ### Limitations:
    - Relies on open ports and banner grabbing
    - NVD database response latency may vary
    - Google search results may require manual verification
    - Not a replacement for enterprise-grade vulnerability scanners
    """)